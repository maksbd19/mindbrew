"""Convert outcome report markdown to PDF and Word documents."""

from __future__ import annotations

import html
import io
import re

import markdown
from docx import Document
from docx.shared import Inches, Pt
from xhtml2pdf import pisa

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_SAFE_FILENAME_RE = re.compile(r"[^\w\s-]", re.UNICODE)
_WHITESPACE_RE = re.compile(r"\s+")
_TABLE_SEPARATOR_RE = re.compile(r"^:?-+:?$")


def safe_filename(title: str, fallback: str = "brewmind-report") -> str:
    cleaned = _SAFE_FILENAME_RE.sub("", title).strip()
    cleaned = _WHITESPACE_RE.sub("-", cleaned).strip("-")
    return cleaned[:80] or fallback


_PDF_STYLES = """
@page {
  size: letter;
  margin: 0.85in 0.85in 1in 0.85in;
  @frame footer {
    -pdf-frame-content: footerContent;
    bottom: 0.45in;
    margin-left: 0.85in;
    margin-right: 0.85in;
    height: 0.35in;
  }
}
body {
  font-family: Helvetica, Arial, sans-serif;
  font-size: 11pt;
  line-height: 1.45;
  color: #111111;
}
.report-header {
  margin-bottom: 20pt;
  padding-bottom: 10pt;
  border-bottom: 1px solid #cccccc;
}
.report-title {
  font-size: 18pt;
  line-height: 1.25;
  margin: 0 0 4pt 0;
  font-weight: bold;
  color: #111111;
}
.report-subtitle {
  font-size: 10pt;
  margin: 0;
  color: #444444;
}
.report-section {
  margin-bottom: 18pt;
}
.section-executive {
  margin-bottom: 22pt;
}
.section-proposal {
  margin-bottom: 16pt;
}
.section-appendix {
  margin-top: 8pt;
  padding-top: 12pt;
  border-top: 1px solid #cccccc;
}
.report-section h2 {
  font-size: 13pt;
  margin: 0 0 8pt 0;
  padding-bottom: 4pt;
  border-bottom: 1px solid #bbbbbb;
  font-weight: bold;
  color: #111111;
}
h3 {
  font-size: 11pt;
  margin: 14pt 0 10pt 0;
  font-weight: bold;
  color: #111111;
}
h4 {
  font-size: 10.5pt;
  margin: 8pt 0 4pt 0;
  font-weight: bold;
  color: #111111;
}
p {
  margin: 0 0 8pt 0;
}
ul, ol {
  margin: 0 0 10pt 0;
  padding-left: 20pt;
}
li {
  margin-bottom: 3pt;
}
strong {
  font-weight: bold;
}
em {
  font-style: italic;
}
hr.section-divider {
  border: none;
  border-top: 1px solid #bbbbbb;
  margin: 16pt 0 18pt 0;
  height: 0;
}
table.data-table {
  width: 100%;
  table-layout: fixed;
  border-collapse: collapse;
  margin: 10pt 0 20pt 0;
}
table.data-table th {
  font-weight: bold;
  font-size: 10pt;
  padding: 9pt 12pt;
  border: 1px solid #999999;
  text-align: left;
  color: #111111;
}
table.data-table td {
  padding: 9pt 12pt;
  border: 1px solid #bbbbbb;
  font-size: 10pt;
  line-height: 1.4;
  vertical-align: top;
  color: #111111;
}
table.data-table th:first-child,
table.data-table td:first-child {
  font-weight: bold;
  width: 30%;
}
table.footer-table {
  width: 100%;
  border: none;
  border-collapse: collapse;
}
table.footer-table td {
  border: none;
  padding: 0;
  font-size: 9pt;
  color: #666666;
}
.footer-right {
  text-align: right;
}
#footerContent {
  border-top: 1px solid #cccccc;
  padding-top: 4pt;
}
"""


def _strip_html_tags(text: str) -> str:
    return re.sub(r"<[^>]+>", "", text).strip()


def _classify_section(heading_text: str) -> str:
    text = _strip_html_tags(heading_text)
    if text == "Executive Summary":
        return "section-executive"
    if text == "Appendix":
        return "section-appendix"
    if re.match(r"^\d+\.", text):
        return "section-proposal"
    return "section-default"


def _enrich_report_html(body_html: str) -> str:
    body_html = re.sub(
        r"<table>",
        '<table class="data-table" width="100%" cellspacing="0" cellpadding="0">',
        body_html,
    )
    body_html = re.sub(r"<hr\s*/?>", '<hr class="section-divider"/>', body_html)

    title_match = re.search(r"<h1>(.*?)</h1>", body_html, re.DOTALL | re.IGNORECASE)
    header_html = ""
    remainder = body_html
    if title_match:
        title_text = html.escape(html.unescape(_strip_html_tags(title_match.group(1))))
        header_html = f"""<div class="report-header">
  <div class="report-title">{title_text}</div>
  <p class="report-subtitle">CRO-ready R&amp;D proposal</p>
</div>"""
        remainder = body_html[title_match.end() :]

    sections = re.split(r"(?=<h2>)", remainder)
    section_html: list[str] = []
    for section in sections:
        if not section.strip():
            continue
        h2_match = re.match(r"(<h2>(.*?)</h2>)(.*)", section, re.DOTALL | re.IGNORECASE)
        if not h2_match:
            section_html.append(section)
            continue
        h2_tag, _h2_inner, content = h2_match.groups()
        css_class = _classify_section(h2_tag)
        section_html.append(
            f'<div class="report-section {css_class}">{h2_tag}{content}</div>'
        )

    footer_html = """<div id="footerContent">
  <table class="footer-table"><tr>
    <td class="footer-right">Page <pdf:pagenumber> of <pdf:pagecount></td>
  </tr></table>
</div>"""

    return header_html + "".join(section_html) + footer_html


def _markdown_to_html(markdown_text: str) -> str:
    body = markdown.markdown(
        markdown_text,
        extensions=["extra", "sane_lists", "nl2br"],
    )
    enriched = _enrich_report_html(body)
    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8"/>
<style>
{_PDF_STYLES}
</style>
</head>
<body>
{enriched}
</body>
</html>"""


def export_report_pdf(markdown_text: str) -> bytes:
    html = _markdown_to_html(markdown_text)
    buffer = io.BytesIO()
    status = pisa.CreatePDF(html, dest=buffer, encoding="utf-8")
    if status.err:
        raise RuntimeError("Failed to generate PDF")
    return buffer.getvalue()


def _add_formatted_runs(paragraph, text: str) -> None:
    parts = _BOLD_RE.split(text)
    for index, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if index % 2 == 1:
            run.bold = True


def _parse_table_row(line: str) -> list[str]:
    inner = line.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    return [cell.strip() for cell in inner.split("|")]


def _is_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def _is_table_separator(cells: list[str]) -> bool:
    meaningful = [cell for cell in cells if cell.strip()]
    if not meaningful:
        return False
    return all(_TABLE_SEPARATOR_RE.match(cell.strip()) for cell in meaningful)


def _configure_docx_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level, size in ((1, 22), (2, 14), (3, 12)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True


def export_report_docx(markdown_text: str) -> bytes:
    doc = Document()
    _configure_docx_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    lines = markdown_text.splitlines()
    paragraph_buffer: list[str] = []
    list_buffer: list[str] = []
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = [_parse_table_row(line) for line in table_buffer]
        table_buffer = []
        if len(rows) < 2:
            for row in rows:
                paragraph_buffer.append(" | ".join(row))
            return

        header = rows[0]
        body_start = 2 if len(rows) > 1 and _is_table_separator(rows[1]) else 1
        body = rows[body_start:]
        if not header:
            return

        table = doc.add_table(rows=1 + len(body), cols=len(header))
        table.style = "Table Grid"
        for col_idx, text in enumerate(header):
            if col_idx < len(table.rows[0].cells):
                _add_formatted_runs(table.rows[0].cells[col_idx].paragraphs[0], text)
        for row_idx, row in enumerate(body):
            for col_idx, text in enumerate(row):
                if col_idx < len(table.rows[row_idx + 1].cells):
                    _add_formatted_runs(
                        table.rows[row_idx + 1].cells[col_idx].paragraphs[0],
                        text,
                    )

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
        paragraph_buffer = []
        if not text:
            return
        paragraph = doc.add_paragraph()
        _add_formatted_runs(paragraph, text)

    def flush_list() -> None:
        nonlocal list_buffer
        if not list_buffer:
            return
        for item in list_buffer:
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(paragraph, item)
        list_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_list()
            flush_paragraph()
            flush_table()
            continue

        if _is_table_row(stripped):
            flush_list()
            flush_paragraph()
            table_buffer.append(stripped)
            continue

        if table_buffer:
            flush_table()

        if stripped.startswith("# "):
            flush_list()
            flush_paragraph()
            doc.add_heading(stripped[2:].strip(), level=1)
            continue

        if stripped.startswith("## "):
            flush_list()
            flush_paragraph()
            doc.add_heading(stripped[3:].strip(), level=2)
            continue

        if stripped.startswith("### "):
            flush_list()
            flush_paragraph()
            doc.add_heading(stripped[4:].strip(), level=3)
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            list_buffer.append(stripped[2:].strip())
            continue

        if line.startswith("  ") and list_buffer:
            list_buffer[-1] = f"{list_buffer[-1]} {stripped}"
            continue

        flush_list()
        paragraph_buffer.append(stripped)

    flush_list()
    flush_paragraph()
    flush_table()

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_report_markdown_preview(markdown_text: str) -> str:
    """Return HTML preview used by tests and optional future endpoints."""
    return _markdown_to_html(markdown_text)
